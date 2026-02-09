import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Join all the xml elements together add add the necessary text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def process_bold(paragraph, text):
    # Split by ** to find bold parts
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def create_docx(md_path, docx_path):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(10)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    table_mode = False
    table_lines = []
    mermaid_mode = False
    
    for line in lines:
        line = line.strip()
        
        # Skip mermaid block markers
        if line.startswith('```mermaid') or (mermaid_mode and line.startswith('```')):
            mermaid_mode = not mermaid_mode
            continue
        if mermaid_mode:
            p = doc.add_paragraph(line)
            p.style = 'Quote' # Use Quote style for code/diagram text
            continue

        # Handle Tables
        if '|' in line:
            if not table_mode:
                # Check if it's a valid table line (starts and ends with | or contains multiple |)
                if line.count('|') >= 2:
                    table_mode = True
                    table_lines = [line]
                else:
                    # Just text with pipe?
                    p = doc.add_paragraph()
                    process_bold(p, line)
            else:
                table_lines.append(line)
            continue
        else:
            if table_mode:
                # End of table, process it
                process_table(doc, table_lines)
                table_mode = False
                table_lines = []
        
        # Handle Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            
        # Handle List Items
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            process_bold(p, line[2:])
            
        # Handle Blockquotes/Notes
        elif line.startswith('>'):
            content = line.replace('>', '').strip()
            if content:
                p = doc.add_paragraph(style='Quote')
                process_bold(p, content)
        
        # Handle Horizontal Rule
        elif line.startswith('---'):
            doc.add_page_break()
            
        # Handle Normal Text
        elif line:
            p = doc.add_paragraph()
            process_bold(p, line)

    # If file ends with table
    if table_mode:
        process_table(doc, table_lines)
        
    doc.save(docx_path)
    print(f"Document saved to {docx_path}")

def process_table(doc, lines):
    # Filter out divider lines (e.g. |---|---|)
    data_lines = [l for l in lines if not set(l.replace('|', '').strip()) <= set('-: ')]
    
    if not data_lines:
        return

    # Determine columns from first line
    headers = [c.strip() for c in data_lines[0].strip('|').split('|')]
    rows = len(data_lines)
    cols = len(headers)
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    for i, line in enumerate(data_lines):
        cells = [c.strip() for c in line.strip('|').split('|')]
        # Handle case where split results in different length than headers
        # Adjust to match cols
        if len(cells) != cols:
            # Simple padding or truncation
            cells = cells[:cols] + [''] * (cols - len(cells))
            
        row_cells = table.rows[i].cells
        for j, cell_text in enumerate(cells):
            # rudimentary bold processing for cells
            p = row_cells[j].paragraphs[0]
            process_bold(p, cell_text)
            
            # Make header row bold/shaded? (Optional, skipping for simplicity)

if __name__ == "__main__":
    base_dir = "/Users/boyoungkim/Desktop/launch-flow/flows/FANC"
    md_file = os.path.join(base_dir, "feasibility_260208_v1.md")
    docx_file = os.path.join(base_dir, "feasibility_260208_v1.docx")
    
    if os.path.exists(md_file):
        create_docx(md_file, docx_file)
    else:
        print(f"File not found: {md_file}")
